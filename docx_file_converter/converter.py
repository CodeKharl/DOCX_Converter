from pathlib import Path
from typing import List

from docx import Document as create_docx
from docx.document import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

DEFAULT_OUTPUT_DOCX = "output.docx"
DEFAULT_FONT_SIZE = 12
DEFAULT_FONT_NAME = "Times New Roman"


def code_to_docx(
    input_files: List[str], output_docx: str, font_name: str, font_size: float
) -> None:
    document: Document = create_docx()

    for file_path in input_files:
        path: Path = Path(file_path)

        document.add_heading(path.name, 2)

        code: str = path.read_text(encoding="utf-8")

        paragraph: Paragraph = document.add_paragraph()

        run: Run = paragraph.add_run(code)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        document.add_paragraph("\n")

    document.save(output_docx)
    print(f"Saved to docx={output_docx}")
